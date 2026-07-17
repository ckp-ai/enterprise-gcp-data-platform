#!/usr/bin/env python3
"""
Apply the Volume 03 related-volume numbering correction.

Usage:
    python apply_volume03_numbering_patch.py \
        Volume_03_Enterprise_Data_Platform_Detailed.docx \
        Volume_03_Enterprise_Data_Platform_Detailed_v1.0.1.docx

The script:
- replaces the stale Related volumes row in Document Control;
- normalizes stale in-line references where found;
- preserves all other document content.
"""

from pathlib import Path
import argparse
from docx import Document

CORRECTED_RELATED_VOLUMES = (
    "Volume 01 Executive Architecture Guide; "
    "Volume 02 Google Cloud Landing Zone; "
    "Volume 04 Data Engineering Pipelines; "
    "Volume 05 Enterprise Data Mesh; "
    "Volume 08 Security and Zero Trust; "
    "Volume 11 Monitoring, SRE and Operations"
)

REPLACEMENTS = {
    "Volume 04 Data Engineering": "Volume 04 Data Engineering Pipelines",
    "Volume 05 Data Mesh": "Volume 05 Enterprise Data Mesh",
    "Volume 07 Security Blueprint": "Volume 08 Security and Zero Trust",
    "Volume 07 Security": "Volume 08 Security and Zero Trust",
    "Volume 10 Operations SRE": "Volume 11 Monitoring, SRE and Operations",
    "Volume 10 Operations": "Volume 11 Monitoring, SRE and Operations",
}

def replace_in_paragraph(paragraph):
    changed = 0
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for old, new in REPLACEMENTS.items():
        new_text = new_text.replace(old, new)
    if new_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = new_text
        changed += 1
    return changed

def patch_docx(source: Path, target: Path):
    doc = Document(source)
    changes = 0
    related_row_found = False

    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            label = row.cells[0].text.strip().lower()
            if label == "related volumes" and len(row.cells) >= 2:
                row.cells[1].text = CORRECTED_RELATED_VOLUMES
                related_row_found = True
                changes += 1

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changes += replace_in_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        changes += replace_in_paragraph(paragraph)

    if not related_row_found:
        raise RuntimeError(
            "Could not find the 'Related volumes' row. "
            "Verify that the source is the Volume 03 detailed DOCX."
        )

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return changes

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source_docx", type=Path)
    parser.add_argument("target_docx", type=Path)
    args = parser.parse_args()

    if not args.source_docx.exists():
        raise FileNotFoundError(args.source_docx)

    changes = patch_docx(args.source_docx, args.target_docx)
    print(f"Patched {changes} location(s).")
    print(f"Saved: {args.target_docx}")

if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""
Apply the Volume 11 stale-reference correction.

Usage:
    python apply_volume11_status_patch.py \
        Volume_11_Monitoring_SRE_Operations_Detailed.docx \
        Volume_11_Monitoring_SRE_Operations_Detailed_v1.0.1.docx
"""

from pathlib import Path
import argparse
from docx import Document

REPLACEMENTS = {
    "Volume 9 DevSecOps has not yet been generated in this sequence. Observability-as-code and release-gate designs are defined here but remain implementation dependencies on Volume 9 and Volume 13.":
    "Volumes 9 and 13 are complete. Volume 9 defines DevSecOps, release-gate and observability-as-code promotion controls; Volume 13 provides the Terraform golden modules, state, policy and apply pipelines required to implement those controls.",

    "Volume 9 DevSecOps - pending":
    "Volume 9 DevSecOps and CI/CD",

    "Implement observability-as-code pipelines as part of the outstanding Volume 9 and Volume 13 work.":
    "Implement and operationalize the observability-as-code pipelines defined by completed Volumes 9 and 13, including promotion, policy validation, release annotation, rollback evidence and golden monitoring/logging modules."
}

def replace_in_paragraph(paragraph):
    original = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    updated = original
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)

    if updated == original:
        return 0

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = updated
    return 1

def patch_docx(source: Path, target: Path):
    doc = Document(source)
    changes = 0

    for paragraph in doc.paragraphs:
        changes += replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changes += replace_in_paragraph(paragraph)

    if changes < 3:
        raise RuntimeError(
            f"Expected at least 3 stale-reference corrections; applied {changes}. "
            "Verify that the source is the original Volume 11 detailed DOCX."
        )

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    return changes

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source_docx", type=Path)
    parser.add_argument("target_docx", type=Path)
    args = parser.parse_args()

    if not args.source_docx.exists():
        raise FileNotFoundError(args.source_docx)

    changes = patch_docx(args.source_docx, args.target_docx)
    print(f"Patched {changes} location(s).")
    print(f"Saved: {args.target_docx}")

if __name__ == "__main__":
    main()
